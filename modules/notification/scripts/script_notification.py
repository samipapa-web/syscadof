# \modules\notification\scripts\script_notification.py
#-------------------------------------------------------
import os
import io
from flask import current_app, session
import pandas as pd
from datetime import datetime
from docx import Document
from docx.enum.text import WD_BREAK
from copy import deepcopy
from docxcompose.composer import Composer
import subprocess
import platform
import time

# ===============================
# FONCTIONS DE FORMATAGE UTILES
# ===============================
def format_nombre(valeur):
    try:
        if pd.isna(valeur):
            return ""
        return "{:,.0f}".format(float(valeur)).replace(",", " ")
    except Exception:
        return str(valeur)

def format_taux(valeur):
    try:
        if pd.isna(valeur) or valeur == "":
            return ""
        return "{:,.2f}".format(float(valeur)).replace(",", " ").replace(".", ",")
    except Exception:
        return str(valeur)

# ===============================
# INSERT HISTORIQUE PostgreSQL
# ===============================
# ===============================
# INSERT HISTORIQUE PostgreSQL
# ===============================
from datetime import datetime

def save_historique(conn, notification, nom_fichier, type_notification, utilisateur, date_creation):
    try:
        conn = conn
        cur = conn.cursor()
        
        cur.execute(
            """
            INSERT INTO notification (
                notification,
                nom_fichier,
                type_notification,
                utilisateur,
                date_creation
            ) VALUES (%s, %s, %s, %s, %s)
            """,
            (
                notification,
                nom_fichier,
                type_notification,
                utilisateur,
                datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            )
        )

        conn.commit()
        cur.close()

    except Exception as e:
        print("Erreur insertion PostgreSQL :", e)

# ===============================
# SCRIPT PRINCIPAL
# ===============================
def run_notif(
        fichier,
        niu,
        centre,
        cri,
        type_personne,
        type_notification,
        modele,
        sortie_dir,
        conn,
        utilisateur
):
    start_time = time.time()

    BASE_DIR = current_app.config["BASE_DIR"]

    # 1️⃣ Chargement fichier Excel
    input_path = os.path.join(sortie_dir, fichier)
    if not os.path.exists(input_path):
        return {"message": "Fichier introuvable."}

    df = pd.read_excel(input_path, engine="openpyxl")
    df.columns = df.columns.str.strip()
    if df.empty:
        return {"message": "Fichier Excel vide."}

    # 2️⃣ Filtrage
    if niu:
        df = df[df["NIU"] == niu]
    else:
        if type_personne:
            df = df[df["NIU"].astype(str).str[0] == type_personne]
        if cri:
            df = df[df["ACRI"] == cri]
        if centre:
            df = df[df["CENTRE"] == centre]

    if df.empty:
        return {"message": "Aucune donnée trouvée après filtrage."}

    # 3️⃣ Nom Excel filtré
    nouveau_nom = fichier.replace("Output_", "_")
    nom_excel_filtre = os.path.splitext(nouveau_nom)[0] + ".xlsx"

    # 4️⃣ Filtre dynamique
    elements_filtre = [str(x) for x in [niu, type_personne, cri, centre] if x]
    filtre = "_".join(elements_filtre) if elements_filtre else "GLOBAL"
    for c in ['<', '>', ':', '"', '/', '\\', '|', '?']:
        filtre = filtre.replace(c, "-")
    filtre = filtre.replace(" ", "_")

    # 5️⃣ Chargement modèle Word
    modele = modele.strip()
    if not modele.lower().endswith(".docx"):
        modele += ".docx"
    #template_path = os.path.join(BASE_DIR, "data", "Templates", modele)
    template_path = os.path.join(current_app.config.get("TEMPLATES_DIR"), modele)
    if not os.path.exists(template_path):
        return {"message": f"❌ Modèle introuvable : {template_path}"}

    base_doc = Document(template_path)

    # 6️⃣ Fonctions de remplacement
    def remplacer_bloc(paragraphes, mapping):
        for p in paragraphes:
            full_text = "".join(run.text for run in p.runs)
            if not full_text: continue
            for key, value in mapping.items():
                full_text = full_text.replace(key, str(value))
            if p.runs:
                p.runs[0].text = full_text
                for run in p.runs[1:]: run.text = ""

    def remplacer_doc(doc, mapping):
        remplacer_bloc(doc.paragraphs, mapping)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    remplacer_bloc(cell.paragraphs, mapping)
        for section in doc.sections:
            remplacer_bloc(section.header.paragraphs, mapping)
            remplacer_bloc(section.footer.paragraphs, mapping)
            if section.first_page_header:
                remplacer_bloc(section.first_page_header.paragraphs, mapping)
            if section.first_page_footer:
                remplacer_bloc(section.first_page_footer.paragraphs, mapping)
            if section.even_page_header:
                remplacer_bloc(section.even_page_header.paragraphs, mapping)
            if section.even_page_footer:
                remplacer_bloc(section.even_page_footer.paragraphs, mapping)

    # 7️⃣ Génération bulletins individuels en mémoire
    date_ref = datetime.now().strftime("%Y%m%d")
    bulletins_mem = []
    for index, row in df.head(100).iterrows():
    #for index, row in df.iterrows():
        doc = deepcopy(base_doc)
        autre_taxe = str(row.get("AUTRE_TAXE") or "").strip()
        mapping = {
            "{{NOTIF}}": type_notification,
            "{{REF_NOTIF}}": f"UTTAD-{index+1}-{date_ref}",
            "{{DIRECTEUR}}": row.get("DIRECTEUR", ""),
            "{{CRIF}}": row.get("CRIF", ""),
            "{{CDIF}}": row.get("CDIF", ""),
            "{{CRIA}}": row.get("CRIA", ""),
            "{{CDIA}}": row.get("CDIA", ""),
            "{{ACRI}}": row.get("ACRI", ""),
            "{{ACDI}}": row.get("ACDI", ""),
            "{{NIU}}": row.get("NIU", ""),
            "{{NOM_RAISON}}": row.get("RAISON_SOCIALE", ""),
            "{{SIGLE}}": row.get("SIGLE", ""),
            "{{TELEPHONE}}": format_nombre(row.get("TELEPHONE")),
            "{{ACTIVITE}}": row.get("ACTIVITE", ""),
            "{{TITRE}}": row.get("TITRE", ""),
            "{{APPEL}}": row.get("APPEL", ""),
            "{{VILLE}}": row.get("VILLE", ""),
            "{{AXE}}": row.get("AXE", ""),
            "{{MATIERE}}": row.get("MATIERE", ""),
            "{{VAL_DECLA}}": format_nombre(row.get("VAL_DECLA")),
            "{{VAL_RECOUP}}": format_nombre(row.get("VAL_RECOUP")),
            "{{ECART_NOTIF}}": format_nombre(row.get("ECART_NOTIF")),
            "{{MARGE_TAUX}}": format_nombre(row.get("MARGE_TAUX")),
            "{{MARGE}}": format_nombre(row.get("MARGE")),
            "{{BASE}}": format_nombre(row.get("BASE")),
            "{{PERIODE}}": row.get("PERIODE", ""),
            "{{LISTE_IMPOTS}}": row.get("LISTE_IMPOTS", ""),
            "{{AUTAX}}": row.get("AUTRE_TAXE", ""),
            "{{PROVENANCE_A}}": row.get("PROVENANCE_A", ""),
            "{{PROVENANCE_B}}": row.get("PROVENANCE_B", ""),
            "{{SOURCE_A}}": row.get("SOURCE_A", ""),
            "{{SOURCE_B}}": row.get("SOURCE_B", ""),
            "{{PERIODE_A}}": row.get("PERIODE_A", ""),
            "{{PERIODE_B}}": row.get("PERIODE_B", ""),
            "{{IS_BASE}}": format_nombre(row.get("IS_BASE")),
            "{{IS_TAUX}}": format_taux(row.get("IS_TAUX")),
            "{{IS_PRINCI}}": format_nombre(row.get("IS_PRINCI")),
            "{{IS_PENAL}}": format_nombre(row.get("IS_PENAL")),
            "{{IS_TOTAL}}": format_nombre(row.get("IS_TOTAL")),
            "{{IRPP_BASE}}": format_nombre(row.get("IRPP_BASE")),
            "{{IRPP_TAUX}}": format_taux(row.get("IRPP_TAUX")),
            "{{IRPP_PRINCI}}": format_nombre(row.get("IRPP_PRINCI")),
            "{{IRPP_PENAL}}": format_nombre(row.get("IRPP_PENAL")),
            "{{IRPP_TOTAL}}": format_nombre(row.get("IRPP_TOTAL")), 
            "{{IRCM_BASE}}": format_nombre(row.get("IRCM_BASE")),
            "{{IRCM_TAUX}}": format_taux(row.get("IRCM_TAUX")),
            "{{IRCM_PRINCI}}": format_nombre(row.get("IRCM_PRINCI")),
            "{{IRCM_PENAL}}": format_nombre(row.get("IRCM_PENAL")),
            "{{IRCM_TOTAL}}": format_nombre(row.get("IRCM_TOTAL")),
            "{{TVA_BASE}}": format_nombre(row.get("TVA_BASE")),
            "{{TVA_TAUX}}": format_taux(row.get("TVA_TAUX")),
            "{{TVA_PRINCI}}": format_nombre(row.get("TVA_PRINCI")),
            "{{TVA_PENAL}}": format_nombre(row.get("TVA_PENAL")),
            "{{TVA_TOTAL}}": format_nombre(row.get("TVA_TOTAL")),
            "{{TAX_PRINCI}}": format_nombre(row.get("TAX_PRINCI")),
            "{{TAX_PENAL}}": format_nombre(row.get("TAX_PENAL")),
            "{{TAX_TOTAL}}": format_nombre(row.get("TAX_TOTAL")),
            "{{AUT_BASE}}": format_nombre(row.get(f"{autre_taxe}_BASE")),
            "{{AUT_TAUX}}": format_taux(row.get(f"{autre_taxe}_TAUX")),
            "{{AUT_PRINCI}}": format_nombre(row.get(f"{autre_taxe}_PRINCI")),
            "{{AUT_PENAL}}": format_nombre(row.get(f"{autre_taxe}_PENAL")),
            "{{AUT_TOTAL}}": format_nombre(row.get(f"{autre_taxe}_TOTAL")),
        }
        remplacer_doc(doc, mapping)

        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        bulletins_mem.append(mem_file)

    print(f"[INFO] {len(bulletins_mem)} bulletins générés en mémoire.")

    # 8️⃣ Fusion Word finale
    
    prefix_map = {
    "Bulletin d’information": "BI",
    "Contrôle sur pièces": "CSP",
    "Déclaration pré-remplie": "DPR",
    "Demande d’éclaircissement": "DE",
    "Mise en demeure": "MED",
    "Listing": "LIST",
    }
    prefix = prefix_map.get(type_notification, "Notif")
    nom_doc = f"{prefix}{os.path.splitext(nouveau_nom)[0]}_{filtre}.docx"
    
    nom_docx_final = os.path.join(sortie_dir, nom_doc)

    if bulletins_mem:
        master_doc = Document(bulletins_mem[0])
        composer = Composer(master_doc)

        for mem_file in bulletins_mem[1:]:
            # 🔹 Ajouter un saut de page
            master_doc.add_page_break()

            # 🔹 Ajouter le document suivant
            composer.append(Document(mem_file))

        composer.save(nom_docx_final)
        print(f"[INFO] Fusion Word finale sauvegardée : {nom_docx_final}")

    # 9️⃣ Conversion PDF
    nom_pdf = nom_doc
    #nom_pdf = nom_doc.replace(".docx", ".pdf")
    #chemin_pdf = os.path.join(sortie_dir, nom_pdf)
    #try:
    #    system_os = platform.system()
    #    libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe" if system_os == "Windows" else "soffice"
    #    subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', sortie_dir, nom_docx_final],
    #                   capture_output=True, text=True)
    #    if not os.path.exists(chemin_pdf):
    #        return {"message": "La conversion PDF a échoué."}
    #    print(f"[INFO] Conversion PDF réussie : {chemin_pdf}")
    #except Exception as e:
    #    return {"message": f"Exception lors de la conversion PDF : {e}"}

    # 9️⃣ Enregistrement historique PostgreSQL
    
    save_historique(
        conn=conn,
        notification=f"Notification {nom_excel_filtre}",
        nom_fichier=nom_pdf,
        type_notification=type_notification,
        utilisateur=utilisateur,
        date_creation = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    )

    end_time = time.time()
    duree = end_time - start_time
    print(f"[INFO] Script terminé avec succès en {duree:.2f} secondes.")

    return {"message": f"{len(df)} bulletins générés avec succès en {duree:.2f} secondes.", "pdf_final": nom_pdf}
