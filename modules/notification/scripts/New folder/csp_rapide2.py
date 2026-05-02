import os
import pandas as pd
from datetime import datetime
import sqlite3
from docx import Document
from docxcompose.composer import Composer
import subprocess
import platform
import time
from multiprocessing import Pool
import tempfile

# ===============================
# Fonctions utilitaires
# ===============================
def format_nombre(valeur):
    try:
        if pd.isna(valeur):
            return ""
        return "{:,.0f}".format(float(valeur)).replace(",", " ")
    except Exception:
        return str(valeur)

def format_taux(valeur):
    try:
        if pd.isna(valeur) or valeur == "":
            return ""
        return "{:,.2f}".format(float(valeur)).replace(",", " ").replace(".", ",")
    except Exception:
        return str(valeur)

# ===============================
# Fonction de génération d’un bulletin (pour multiprocessing)
# ===============================
def generer_bulletin(row_data):
    row, template_path, script_name, date_ref = row_data
    doc = Document(template_path)
    
    mapping = {
        "{{NOTIF}}": script_name,
        "{{REF_NOTIF}}": f"UTTAD-{row.name+1}-{date_ref}",
        "{{DIRECTEUR}}": row.get("DIRECTEUR", ""),
        "{{CRIF}}": row.get("CRIF", ""),
        "{{CDIF}}": row.get("CDIF", ""),
        "{{CRIA}}": row.get("CRIA", ""),
        "{{CDIA}}": row.get("CDIA", ""),
        "{{ACRI}}": row.get("ACRI", ""),
        "{{ACDI}}": row.get("ACDI", ""),
        "{{NIU}}": row.get("NIU", ""),
        "{{NOM_RAISON}}": row.get("NOM_RAISON", ""),
        "{{SIGLE}}": row.get("SIGLE", ""),
        "{{TELEPHONE}}": format_nombre(row.get("TELEPHONE")),
        "{{ACTIVITE}}": row.get("ACTIVITE", ""),
        "{{TITRE}}": row.get("TITRE", ""),
        "{{APPEL}}": row.get("APPEL", ""),
        "{{VILLE}}": row.get("VILLE", ""),
        "{{AXE}}": row.get("AXE", ""),
        "{{MATIERE}}": row.get("MATIERE", ""),
        "{{VAL_DECLA}}": format_nombre(row.get("VAL_DECLA")),
        "{{VAL_RECOUP}}": format_nombre(row.get("VAL_RECOUP")),
        "{{ECART_NOTIF}}": format_nombre(row.get("ECART_NOTIF")),
        "{{MARGE_TAUX}}": format_nombre(row.get("MARGE_TAUX")),
        "{{MARGE}}": format_nombre(row.get("MARGE")),
        "{{BASE}}": format_nombre(row.get("BASE")),
        "{{PERIODE}}": row.get("PERIODE", ""),
        "{{LISTE_IMPOTS}}": row.get("LISTE_IMPOTS", ""),
        "{{AUTAX}}": row.get("AUTRE_TAXE", ""),
        "{{PROVENANCE_A}}": row.get("PROVENANCE_A", ""),
        "{{PROVENANCE_B}}": row.get("PROVENANCE_B", ""),
        "{{SOURCE_A}}": row.get("SOURCE_A", ""),
        "{{SOURCE_B}}": row.get("SOURCE_B", ""),
        "{{PERIODE_A}}": row.get("PERIODE_A", ""),
        "{{PERIODE_B}}": row.get("PERIODE_B", ""),
        "{{IS_BASE}}": format_nombre(row.get("IS_BASE")),
        "{{IS_TAUX}}": format_taux(row.get("IS_TAUX")),
        "{{IS_PRINCI}}": format_nombre(row.get("IS_PRINCI")),
        "{{IS_PENAL}}": format_nombre(row.get("IS_PENAL")),
        "{{IS_TOTAL}}": format_nombre(row.get("IS_TOTAL")),
        "{{IRPP_BASE}}": format_nombre(row.get("IRPP_BASE")),
        "{{IRPP_TAUX}}": format_taux(row.get("IRPP_TAUX")),
        "{{IRPP_PRINCI}}": format_nombre(row.get("IRPP_PRINCI")),
        "{{IRPP_PENAL}}": format_nombre(row.get("IRPP_PENAL")),
        "{{IRPP_TOTAL}}": format_nombre(row.get("IRPP_TOTAL")), 
        "{{IRCM_BASE}}": format_nombre(row.get("IRCM_BASE")),
        "{{IRCM_TAUX}}": format_taux(row.get("IRCM_TAUX")),
        "{{IRCM_PRINCI}}": format_nombre(row.get("IRCM_PRINCI")),
        "{{IRCM_PENAL}}": format_nombre(row.get("IRCM_PENAL")),
        "{{IRCM_TOTAL}}": format_nombre(row.get("IRCM_TOTAL")),
        "{{TVA_BASE}}": format_nombre(row.get("TVA_BASE")),
        "{{TVA_TAUX}}": format_taux(row.get("TVA_TAUX")),
        "{{TVA_PRINCI}}": format_nombre(row.get("TVA_PRINCI")),
        "{{TVA_PENAL}}": format_nombre(row.get("TVA_PENAL")),
        "{{TVA_TOTAL}}": format_nombre(row.get("TVA_TOTAL")),
        "{{TAX_PRINCI}}": format_nombre(row.get("TAX_PRINCI")),
        "{{TAX_PENAL}}": format_nombre(row.get("TAX_PENAL")),
        "{{TAX_TOTAL}}": format_nombre(row.get("TAX_TOTAL")),
        "{{AUT_BASE}}": format_nombre(row.get(f"{row.get('AUTRE_TAXE','')}_BASE")),
        "{{AUT_TAUX}}": format_taux(row.get(f"{row.get('AUTRE_TAXE','')}_TAUX")),
        "{{AUT_PRINCI}}": format_nombre(row.get(f"{row.get('AUTRE_TAXE','')}_PRINCI")),
        "{{AUT_PENAL}}": format_nombre(row.get(f"{row.get('AUTRE_TAXE','')}_PENAL")),
        "{{AUT_TOTAL}}": format_nombre(row.get(f"{row.get('AUTRE_TAXE','')}_TOTAL"))
    }

    # Remplacement texte dans doc
    def remplacer_bloc(paragraphes, mapping):
        for p in paragraphes:
            full_text = "".join(run.text for run in p.runs)
            for key, value in mapping.items():
                full_text = full_text.replace(key, str(value))
            if p.runs:
                p.runs[0].text = full_text
                for run in p.runs[1:]: run.text = ""
    
    remplacer_bloc(doc.paragraphs, mapping)
    for table in doc.tables:
        for row_tbl in table.rows:
            for cell in row_tbl.cells:
                remplacer_bloc(cell.paragraphs, mapping)

    # Sauvegarde dans fichier temporaire
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_file.name)
    return tmp_file.name

# ===============================
# Script principal run_notif
# ===============================
def run_notif(
        fichier, niu, centre, cri, type_personne, type_notification,
        modele, sortie_dir, db_path, script_name, utilisateur,
        nb_processus=4
):
    start_time = time.time()
    print("[INFO] Démarrage du script de notification.")
    print("UTILISATEUR RECU :", utilisateur)

    input_path = os.path.join(sortie_dir, fichier)
    if not os.path.exists(input_path):
        return {"message": "Fichier introuvable."}

    df = pd.read_excel(input_path, engine="openpyxl")
    df.columns = df.columns.str.strip()
    if df.empty:
        return {"message": "Fichier Excel vide."}
    print(f"[INFO] Fichier Excel chargé avec {len(df)} lignes.")

    # Filtrage
    if niu:
        df = df[df["NIU"] == niu]
    else:
        if type_personne: df = df[df["NIU"].astype(str).str[0] == type_personne]
        if cri: df = df[df["ACRI"] == cri]
        if centre: df = df[df["CENTRE"] == centre]
    if df.empty:
        return {"message": "Aucune donnée trouvée après filtrage."}

    nouveau_nom = fichier.replace("Output_", "_")
    nom_excel_filtre = os.path.splitext(nouveau_nom)[0] + ".xlsx"
    date_ref = datetime.now().strftime("%Y%m%d")

    # Filtre dynamique
    elements_filtre = [str(x) for x in [niu, type_personne, cri, centre] if x]
    filtre = "_".join(elements_filtre) if elements_filtre else "GLOBAL"
    for c in ['<', '>', ':', '"', '/', '\\', '|', '?']:
        filtre = filtre.replace(c, "-")
    filtre = filtre.replace(" ", "_")

    template_path = os.path.join("data", "Templates", modele)
    if not os.path.exists(template_path):
        return {"message": f"Modèle introuvable : {template_path}"}

    # Préparer les données pour multiprocessing
    rows_data = [(row, template_path, script_name, date_ref) for _, row in df.iterrows()]

    # Multiprocessing
    with Pool(processes=nb_processus) as pool:
        bulletins_fichiers = pool.map(generer_bulletin, rows_data)

    print(f"[INFO] {len(bulletins_fichiers)} bulletins générés en parallèle.")

    # Fusion finale
    nom_doc = f"BI{os.path.splitext(nouveau_nom)[0]}_{filtre}.docx"
    nom_docx_final = os.path.join(sortie_dir, nom_doc)
    master_doc = Document(bulletins_fichiers[0])
    composer = Composer(master_doc)
    for f in bulletins_fichiers[1:]:
        composer.append(Document(f))
    composer.save(nom_docx_final)
    print(f"[INFO] Fusion Word finale sauvegardée : {nom_docx_final}")

    # Conversion PDF
    nom_pdf = nom_doc.replace(".docx", ".pdf")
    chemin_pdf = os.path.join(sortie_dir, nom_pdf)
    try:
        system_os = platform.system()
        libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe" if system_os == "Windows" else "soffice"
        subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', sortie_dir, nom_docx_final], capture_output=True, text=True)
        if not os.path.exists(chemin_pdf):
            return {"message": "La conversion PDF a échoué."}
        print(f"[INFO] Conversion PDF réussie : {chemin_pdf}")
    except Exception as e:
        return {"message": f"Exception lors de la conversion PDF : {e}"}

    # Nettoyage fichiers temporaires
    for f in bulletins_fichiers:
        try: os.remove(f)
        except Exception: pass

    # Enregistrement base SQLite
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO notification (
            notification, nom_fichier, type_notification,
            script_utilise, utilisateur, date_creation
        )
        VALUES (?, ?, ?, ?, ?, ?)
    """, (f"Notification {nom_excel_filtre}", nom_pdf, type_notification, script_name, utilisateur, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    conn.commit()
    conn.close()

    end_time = time.time()
    duree = end_time - start_time
    print(f"[INFO] Script terminé avec succès en {duree:.2f} secondes.")

    return {"message": f"{len(df)} bulletins générés avec succès en {duree:.2f} secondes.", "pdf_final": nom_pdf}