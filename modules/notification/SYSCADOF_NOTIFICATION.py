#
# MODULE 6: RESTITUTION/NOTIFICATION
################################################################

# \modules\notification\templates\notification.html
#-------------------------------------------------------

<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<title>SYSCADOF - Notification</title>

<script src="https://cdnjs.cloudflare.com/ajax/libs/ace/1.15.0/ace.js"></script>

<style>
/* ============================
   Structure générale
============================ */
html, body { margin:0; padding:0; font-family: Arial, sans-serif; background: #f3f4f6; }

/* Bannière fixée en haut */
.banniere-container {
    position: fixed;
    top:0; left:0; width:100%; z-index:1000;
}
.banniere-container img {
    width:100%; height:auto; display:block;
}

/* Variables dynamiques */
:root { --banner-height:0px; }

/* Sidebar fixe */
.sidebar {
    position: fixed;
    top: var(--banner-height);
    left:0; bottom:0;
    width:110px;
    background-color: #1e293b;
    overflow-y: auto;
    padding-top:10px;
}
.sidebar a {
    display:block;
    margin:10px;
    padding:10px 5px;
    text-align:center;
    color:white;
    text-decoration:none;
    font-weight:bold;
    border-radius:4px;
    background-color: #127291;
    font-size:12px;
}
.sidebar a:hover { background-color:#0d5a72; }

/* Contenu principal */
.main-content {
    margin-left:110px;
    padding:20px;
    padding-top: calc(20px + var(--banner-height));
    overflow-x:auto;
}

/* Formulaire */
h3 { color:#2099b6; }
select, input { padding:6px; margin:5px 0; width:100%; }
button { padding:6px 12px; background:#2099b6; color:white; border:none; cursor:pointer; border-radius:4px; margin-top:5px; }
button:hover { background:#167d94; }
.form-grid { display:flex; flex-direction:column; gap:12px; margin-top:15px; margin-bottom:15px; }
.form-row { display:grid; grid-template-columns:230px 1fr; align-items:center; gap:15px; }
.form-row label { font-weight:bold; }

/* Table */
table { width:100%; border-collapse: collapse; margin-top:20px; background:white; }
th { background:#2099b6; color:white; padding:10px; text-align:center; }
td { padding:8px; border-bottom:1px solid #ddd; text-align:center; }
.separator { border:none; border-top:2px solid #2099b6; margin:40px 0 10px 0; }
#log { margin-top:10px; height:30px; background:white; border:1px solid #ccc; padding:10px; overflow-y:auto; white-space:pre-wrap; font-family: monospace; }
.open-btn { display:inline-block; padding:4px 8px; background-color:#2099b6; color:white; border-radius:4px; text-decoration:none; font-size:0.9rem; }
.open-btn:hover { background-color:#167d94; }
</style>
</head>

<body>

<div class="banniere-container">
    <img src="{{ url_for('static', filename='images/notification.jpg') }}" alt="Notification">
</div>

<aside class="sidebar">
    <a href="/">Accueil</a>

    {% if modules.get("Ingestion") %}<a href="/ingestion">Ingestion</a>{% endif %}
    {% if modules.get("Apurement") %}<a href="/traitement/editor">Apurement</a>{% endif %}
    {% if modules.get("Croisement") %}<a href="/analyse">Croisement</a>{% endif %}
    {% if modules.get("Orientation") %}<a href="/exploitation">Orientation</a>{% endif %}
    {% if modules.get("Valorisation") %}<a href="/valorisation">Valorisation</a>{% endif %}
    {% if modules.get("Restitution") %}<a href="/notification">Restitution</a>{% endif %}
</aside>

<main class="main-content">

    <h3>Paramétrage des notifications</h3>
    <div class="form-grid">
        <div class="form-row"><label for="fichierSelect">Fichier à exploiter</label><select id="fichierSelect"></select></div>
        <div class="form-row"><label for="scriptSelect">Type de notification</label><select id="scriptSelect"></select></div>
        <div class="form-row"><label for="modeleSelect">Modèle de notification</label><select id="modeleSelect"></select></div>
        <div class="form-row"><label for="niuInput">NIU</label><input type="text" id="niuInput" placeholder="Saisir NIU"></div>
        <div class="form-row"><label for="typePersonneSelect">Type d'entité</label><select id="typePersonneSelect"></select></div>
        <div class="form-row"><label for="criSelect">Région fiscale</label><select id="criSelect"></select></div>
        <div class="form-row"><label for="centreSelect">Unité de gestion</label><select id="centreSelect"></select></div>
    </div>

    <button onclick="executeNotification()">Générer les notifications</button>
	<button onclick="ouvrirNouveauModele()">Créer un modèle</button>
    <div id="log"></div>

    <hr class="separator">
    <h3>Historique des notifications générées</h3>
    <table>
        <thead>
            <tr>
                <th>ID</th>
                <th>Fichier généré</th>
                <th>Type notification</th>
                <th>Utilisateur</th>
                <th>Date création</th>
                <th>Ouvrir</th>
            </tr>
        </thead>
        <tbody id="historiqueBody"></tbody>
    </table>

</main>

<script>
// ============================
// Ajustement bannière dynamique
// ============================
function adjustLayout() {
    const banner = document.querySelector(".banniere-container");
    if (banner) {
        const height = banner.offsetHeight + "px";
        document.documentElement.style.setProperty("--banner-height", height);
    }
}
window.addEventListener("load", adjustLayout);
window.addEventListener("resize", adjustLayout);

// ============================
// Initialisation des selects et historique
// ============================
document.addEventListener("DOMContentLoaded", () => {
    initSelects();
    loadHistorique();
    document.getElementById("scriptSelect").addEventListener("change", loadModeles);
});

async function initSelects() {
    try {
        const res = await fetch("/notification/api/init");
        const data = await res.json();

        populateSelect("fichierSelect", data.files);
        populateSelect("centreSelect", data.centres);
        populateSelect("criSelect", data.cris);
        populateSelect("typePersonneSelect", ["P","M"]);
        populateSelect("scriptSelect", data.types_notif);

    } catch(e) { console.error(e); }
}

// ============================
// Chargement dynamique modèles
// ============================
async function loadModeles() {
    const type = document.getElementById("scriptSelect").value;
    if(!type){ populateSelect("modeleSelect", []); return; }
    try {
        const res = await fetch(`/notification/api/modeles?type=${encodeURIComponent(type)}`);
        const data = await res.json();
        populateSelect("modeleSelect", data);
    } catch(e) { console.error(e); }
}

// ============================
// Utils
// ============================
function populateSelect(id, items) {
    const select = document.getElementById(id);
    if(!select) return;
    select.innerHTML = "<option value=''>-- Sélectionner --</option>";
    items.forEach(i => { select.innerHTML += `<option value="${i}">${i}</option>`; });
}

// ============================
// Exécution script
// ============================
async function executeNotification() {
    const payload = {
        fichier: document.getElementById("fichierSelect").value,
        niu: document.getElementById("niuInput").value,
        centre: document.getElementById("centreSelect").value,
        cri: document.getElementById("criSelect").value,
        type_personne: document.getElementById("typePersonneSelect").value,
        type_notification: document.getElementById("scriptSelect").value,
        modele: document.getElementById("modeleSelect").value
    };
    const log = document.getElementById("log");
    log.textContent = "⏳ Exécution en cours...";
    try {
        const res = await fetch("/notification/api/execute", {
            method:"POST", headers:{"Content-Type":"application/json"}, body:JSON.stringify(payload)
        });
        const result = await res.json();
        log.textContent = result.success ? "✅ " + result.message : "❌ " + result.message;
        if(result.success) loadHistorique();
    } catch(e){ log.textContent = "❌ Erreur serveur : " + e.message; }
}

// ============================
// Historique avec ouverture PDF.html
// ============================
async function loadHistorique() {
    try {
        const res = await fetch("/notification/api/historique");
        const data = await res.json();
        const tbody = document.getElementById("historiqueBody");
        tbody.innerHTML = "";
        data.forEach(r => {
            tbody.innerHTML += `
                <tr>
                    <td>${r.id}</td>
                    <td>${r.nom_fichier}</td>
                    <td>${r.type_notification || ''}</td>
                    <td>${r.utilisateur || 'system'}</td>
                    <td>${r.date_creation || ''}</td>
                    <td>
                        <button class="open-btn" onclick="ouvrirPDF('${r.nom_fichier}')">Ouvrir</button>
                    </td>
                </tr>`;
        });
    } catch(e){ console.error(e); }
}

// ============================
// Ouvrir PDF via PDF.html
// ============================
function ouvrirPDF(nomFichier) {
    window.open(`/notification/pdf/${encodeURIComponent(nomFichier)}`, "_blank");
}

// ============================
// Ouvrir Modele
// ============================
function ouvrirNouveauModele() {
    let modele = document.getElementById("modeleSelect").value;

    if (!modele) {
        alert("Veuillez sélectionner un modèle de notification.");
        return;
    }

    // ✅ Ajouter automatiquement .docx si absent
    if (!modele.endsWith(".docx")) {
        modele += ".docx";
    }

    // ✅ Redirection avec le bon nom
    // window.open(`/notification/doc?modele=${encodeURIComponent(modele)}`, "_blank");
	window.location.href = `/notification/doc?modele=${encodeURIComponent(modele)}`;
}

</script>

</body>
</html>


# \modules\notification\routes.py
#-------------------------------------------------------
import os
import importlib
import psycopg2
from psycopg2.extras import RealDictCursor
from flask import Blueprint, render_template, jsonify, request, current_app, send_from_directory, session
from werkzeug.utils import secure_filename
from .services import get_output_files, get_types_notification, get_filtres_data, get_modeles_par_type

modules = {
    "Ingestion": True,
    "Apurement": True,
    "Croisement": True,
    "Orientation": True,
    "Valorisation": True,
    "Restitution": True
}

notification_bp = Blueprint(
    "notification",
    __name__,
    url_prefix="/notification",
    template_folder="templates",
    static_folder="static"
)

# ===============================
# PAGE PRINCIPALE
# ===============================
@notification_bp.route("/")
def notification_page():
    return render_template("notification.html")

# ===============================
# INITIALISATION
# ===============================
@notification_bp.route("/api/init")
def init_data():
    files = get_output_files()
    types_notif = get_types_notification()
    centres, cris = get_filtres_data()

    return jsonify({
        "files": files,
        "types_notif": types_notif,
        "centres": centres,
        "cris": cris,
        "modeles": []  # modèles chargés dynamiquement via /api/modeles
    })

# ===============================
# CHARGEMENT DES MODÈLES PAR TYPE
# ===============================
@notification_bp.route("/api/modeles")
def modeles_par_type():
    type_notif = request.args.get("type")
    if not type_notif:
        return jsonify([])

    modeles = get_modeles_par_type(type_notif)
    return jsonify(modeles)

# ===============================
# EXECUTION SCRIPT
# ===============================
@notification_bp.route("/api/execute", methods=["POST"])
def execute_script():
    data = request.json

    fichier = data.get("fichier")
    niu = data.get("niu")
    centre = data.get("centre")
    cri = data.get("cri")
    type_personne = data.get("type_personne")
    type_notification = data.get("type_notification")
    modele = data.get("modele")

    script_name = "script_notification"  # nom du script

    utilisateur = session.get("user", "system")

    try:
        # Import dynamique du script
        module_path = f"modules.notification.scripts.{script_name}"
        script_module = importlib.import_module(module_path)

        # Appel de la fonction run_notif du script
        result = script_module.run_notif(
            fichier=fichier,
            niu=niu,
            centre=centre,
            cri=cri,
            type_personne=type_personne,
            type_notification=type_notification,
            modele=modele,
            sortie_dir=current_app.config["SORTIE_DIR"],
            db_url=current_app.config["DB_URL"],  # PostgreSQL
            utilisateur=utilisateur
        )

        return jsonify({
            "success": True,
            "message": result.get("message"),
            "pdf_final": result.get("pdf_final")
        })

    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

# ===============================
# HISTORIQUE
# ===============================
@notification_bp.route("/api/historique")
def historique():
    db_url = current_app.config.get("DB_URL")
    if not db_url:
        return jsonify({"error": "DB_URL non configuré"})

    try:
        conn = psycopg2.connect(db_url, cursor_factory=RealDictCursor)
        cur = conn.cursor()
        cur.execute("SELECT * FROM notification ORDER BY id DESC")
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)})

# ===============================
# PAGE VIEWER PDF
# ===============================
@notification_bp.route("/pdf/<path:nom>")
def pdf_viewer(nom):
    safe_nom = secure_filename(nom)
    return render_template("PDF.html", fichier=safe_nom)

# ===============================
# SERVEUR PDF
# ===============================
@notification_bp.route("/pdf_file/<path:nom>")
def pdf_file(nom):
    sortie_dir = current_app.config.get("SORTIE_DIR")
    if not sortie_dir:
        return "SORTIE_DIR non configuré", 500

    safe_nom = secure_filename(nom)
    chemin_fichier = os.path.join(sortie_dir, safe_nom)
    if not os.path.exists(chemin_fichier):
        return f"Fichier introuvable : {safe_nom}", 404

    return send_from_directory(sortie_dir, safe_nom)

# ===============================
# LISTE ET UPLOAD DES MODELES WORD
# ===============================
@notification_bp.route("/api/templates")
def list_templates():
    folder = os.path.join("data", "templates")
    files = [f for f in os.listdir(folder) if f.endswith(".docx")]
    return jsonify(files)

@notification_bp.route("/template/<filename>")
def get_template(filename):
    folder = os.path.join("data", "templates")
    return send_from_directory(folder, filename)

@notification_bp.route("/doc")
def doc_page():
    return render_template("doc.html", modules=modules)



# \modules\notification\services.py
#-------------------------------------------------------
import os
import pandas as pd

# Répertoire de base du projet
BASE_DIR = os.getcwd()


# ===============================
# Liste des fichiers de sortie (Excel)
# ===============================
def get_output_files():
    sortie_dir = os.path.join(BASE_DIR, "data", "Sortie")
    if not os.path.exists(sortie_dir):
        return []
    return [f for f in os.listdir(sortie_dir) if f.endswith(".xlsx")]

# ===============================
# Liste des fichiers de sortie (Excel) commençant par Output_ et finissant par .xlsx
# ===============================
def get_output_files():
    sortie_dir = os.path.join(BASE_DIR, "data", "Sortie")
    if not os.path.exists(sortie_dir):
        return []
    
    return [
        f for f in os.listdir(sortie_dir)
        if f.startswith("Output_") and f.endswith(".xlsx")
    ]


# ===============================
# Types de notification disponibles
# ===============================
def get_types_notification():
    path = os.path.join(BASE_DIR, "data", "Templates", "NOTIFS.xlsx")
    if not os.path.exists(path):
        return []

    df = pd.read_excel(path)
    df.columns = df.columns.str.strip()

    if "Type_notification" not in df.columns:
        return []

    return sorted(df["Type_notification"].dropna().astype(str).unique().tolist())


# ===============================
# Modèles de notification selon le type
# ===============================
def get_modeles_par_type(type_notif):
    path = os.path.join(BASE_DIR, "data", "Templates", "NOTIFS.xlsx")
    if not os.path.exists(path):
        return []

    df = pd.read_excel(path)
    df.columns = df.columns.str.strip()
    if "Type_notification" not in df.columns or "Modele_notification" not in df.columns:
        return []

    df = df.dropna(subset=["Type_notification", "Modele_notification"])
    modeles = df[df["Type_notification"] == type_notif]["Modele_notification"].astype(str).unique()
    return sorted(modeles.tolist())


# ===============================
# Données pour les filtres : CENTRE et CRI
# ===============================
def get_filtres_data():
    path = os.path.join(BASE_DIR, "data", "Templates", "FILTRES.xlsx")
    if not os.path.exists(path):
        return [], []

    centres, cris = [], []

    try:
        df_centre = pd.read_excel(path, sheet_name="CENTRE")
        if "CENTRE" in df_centre.columns:
            centres = df_centre["CENTRE"].dropna().astype(str).tolist()
    except Exception:
        centres = []

    try:
        df_cri = pd.read_excel(path, sheet_name="CRI")
        if "CRI" in df_cri.columns:
            cris = df_cri["CRI"].dropna().astype(str).tolist()
    except Exception:
        cris = []

    return centres, cris


# \modules\notification\templates\doc.html
#-------------------------------------------------------

<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<title>Gestion des modèles Word</title>

<style>

/* ============================
   Structure générale
============================ */
html, body {
    margin:0;
    padding:0;
    font-family: Arial, sans-serif;
    background:#f3f4f6;
}

/* ============================
   Bannière
============================ */
.banniere-container {
    position: fixed;
    top:0; left:0; width:100%; z-index:1000;
}
.banniere-container img {
    width:100%; height:auto; display:block;
}

/* Variable dynamique */
:root { --banner-height:0px; }

/* ============================
   Sidebar
============================ */
.sidebar {
    position: fixed;
    top: var(--banner-height);
    left:0; bottom:0;
    width:110px;
    background-color: #1e293b;
    overflow-y:auto;
    padding-top:10px;
}
.sidebar a {
    display:block;
    margin:10px;
    padding:10px 5px;
    text-align:center;
    color:white;
    text-decoration:none;
    font-weight:bold;
    border-radius:4px;
    background-color:#127291;
    font-size:12px;
}
.sidebar a:hover { background-color:#0d5a72; }

/* ============================
   Contenu principal
============================ */
.main-content {
    margin-left:110px;
    padding:20px;
    padding-top: calc(20px + var(--banner-height));
}

/* ============================
   Container
============================ */
.container {
    max-width:900px;
    margin:auto;
    background:white;
    padding:20px;
    border-radius:8px;
    box-shadow:0 2px 10px rgba(0,0,0,0.1);
}

/* ============================
   Formulaire
============================ */
h3 { color:#2099b6; }

select, button, input {
    padding:8px;
    margin-top:10px;
    width:100%;
}

button {
    background:#2099b6;
    color:white;
    border:none;
    border-radius:4px;
    cursor:pointer;
}
button:hover { background:#167d94; }

.form-group {
    display:flex;
    flex-direction:column;
    margin-top:10px;
}

/* ============================
   Viewer
============================ */
iframe {
    margin-top:20px;
    width:100%;
    height:600px;
    border:1px solid #ccc;
    background:white;
}

</style>
</head>

<body>

<!-- ============================
     Bannière
============================ -->
<div class="banniere-container">
    <img src="{{ url_for('static', filename='images/modelenotif.jpg') }}" alt="Notification">
</div>

<!-- ============================
     Sidebar
============================ -->
<aside class="sidebar">
    <a href="/">Accueil</a>

    {% if modules.get("Ingestion") %}<a href="/ingestion">Ingestion</a>{% endif %}
    {% if modules.get("Apurement") %}<a href="/traitement/editor">Apurement</a>{% endif %}
    {% if modules.get("Croisement") %}<a href="/analyse">Croisement</a>{% endif %}
    {% if modules.get("Orientation") %}<a href="/exploitation">Orientation</a>{% endif %}
    {% if modules.get("Valorisation") %}<a href="/valorisation">Valorisation</a>{% endif %}
    {% if modules.get("Restitution") %}<a href="/notification">Restitution</a>{% endif %}
</aside>

<!-- ============================
     Contenu principal
============================ -->
<main class="main-content">

<div class="container">

    <!-- ============================
         Modifier modèle
    ============================ -->
    <h3>Modifier un modèle existant</h3>

    <label>Choisir un modèle</label>
    <select id="modeleSelect"></select>

    <button onclick="ouvrirModele()">Ouvrir</button>

    <hr>

    <!-- ============================
         Ajouter modèle
    ============================ -->
    <h3>Ajouter un nouveau modèle</h3>

    <label>Type de notification</label>
    <select id="typeNotifSelect"></select>

    <div class="form-group">
        <label>Sélectionner un fichier (.docx)</label>
        <input type="file" id="fileInput" accept=".docx" onchange="remplirIntitule()">
    </div>

    <div class="form-group">
        <label>Intitulé du modèle</label>
        <input type="text" id="intituleInput" placeholder="Nom du modèle">
    </div>

    <button onclick="stockerModele()">Stocker le modèle</button>

    <div id="uploadLog"></div>

    <!-- ============================
         Viewer
    ============================ -->
    <iframe id="viewer"></iframe>

</div>

</main>

<script>

// ============================
// Ajustement bannière dynamique
// ============================
function adjustLayout() {
    const banner = document.querySelector(".banniere-container");
    if (banner) {
        const height = banner.offsetHeight + "px";
        document.documentElement.style.setProperty("--banner-height", height);
    }
}
window.addEventListener("load", adjustLayout);
window.addEventListener("resize", adjustLayout);

// ============================
// Initialisation
// ============================
document.addEventListener("DOMContentLoaded", async () => {
    await chargerModeles();
    await chargerTypesNotification();

    const params = new URLSearchParams(window.location.search);
    let modele = params.get("modele");

    if (modele) {
        const select = document.getElementById("modeleSelect");
        const options = Array.from(select.options).map(o => o.value);

        if (options.includes(modele)) {
            select.value = modele;
            ouvrirModele();
        }
    }
});

// ============================
// Charger modèles
// ============================
async function chargerModeles() {
    try {
        const res = await fetch("/notification/api/templates");
        const data = await res.json();

        const select = document.getElementById("modeleSelect");
        select.innerHTML = "<option value=''>-- Sélectionner --</option>";

        data.forEach(m => {
            select.innerHTML += `<option value="${m}">${m}</option>`;
        });

    } catch(e) { console.error(e); }
}

// ============================
// Charger types notification
// ============================
async function chargerTypesNotification() {
    try {
        const res = await fetch("/notification/api/types_notification");
        const data = await res.json();

        const select = document.getElementById("typeNotifSelect");
        select.innerHTML = "<option value=''>-- Sélectionner --</option>";

        data.forEach(t => {
            select.innerHTML += `<option value="${t}">${t}</option>`;
        });

    } catch(e) { console.error(e); }
}

// ============================
// Remplir intitulé automatiquement
// ============================
function remplirIntitule() {
    const fileInput = document.getElementById("fileInput");
    const intituleInput = document.getElementById("intituleInput");

    if (fileInput.files.length > 0) {
        let nom = fileInput.files[0].name;
        nom = nom.replace(/\.docx$/i, "");
        intituleInput.value = nom;
    }
}

// ============================
// Stocker modèle
// ============================
async function stockerModele() {
    const fileInput = document.getElementById("fileInput");
    const typeNotif = document.getElementById("typeNotifSelect").value;
    const intitule = document.getElementById("intituleInput").value.trim();
    const log = document.getElementById("uploadLog");

    if (!fileInput.files.length) {
        log.textContent = "❌ Veuillez sélectionner un fichier.";
        return;
    }

    if (!typeNotif) {
        log.textContent = "❌ Veuillez choisir un type.";
        return;
    }

    if (!intitule) {
        log.textContent = "❌ Veuillez saisir un intitulé.";
        return;
    }

    const file = fileInput.files[0];

    const formData = new FormData();
    formData.append("file", file);
    formData.append("type_notification", typeNotif);
    formData.append("intitule", intitule);

    log.textContent = "⏳ Enregistrement...";

    try {
        const res = await fetch("/notification/api/upload_template", {
            method:"POST",
            body:formData
        });

        const result = await res.json();

        log.textContent = result.success
            ? "✅ " + result.message
            : "❌ " + result.message;

        if (result.success) {
            chargerModeles();
        }

    } catch(e) {
        log.textContent = "❌ Erreur serveur : " + e.message;
    }
}

// ============================
// Ouvrir modèle
// ============================
function ouvrirModele() {
    const modele = document.getElementById("modeleSelect").value;

    if (!modele) {
        alert("Veuillez sélectionner un modèle.");
        return;
    }

    document.getElementById("viewer").src =
        `/notification/template/${encodeURIComponent(modele)}`;
}

</script>

</body>
</html>

# \modules\notification\routes_doc.py   
#---------------------------------------

from flask import Blueprint, request, jsonify
import os
import pandas as pd
from werkzeug.utils import secure_filename

# ✅ Blueprint dédié
doc_bp = Blueprint("doc_bp", __name__)

# ===============================
# A. Lire les types depuis Excel
# ===============================
@doc_bp.route("/api/types_notification")
def get_types_notification():
    path = os.path.join("data", "templates", "NOTIFS.xlsx")

    df = pd.read_excel(path, sheet_name="NOTIFS")

    types = (
        df["Type_notification"]
        .dropna()
        .astype(str)
        .str.strip()
        .unique()
        .tolist()
    )

    return jsonify(types)

# ===============================
# B. Upload + enregistrement Excel
# ===============================
@doc_bp.route("/api/upload_template", methods=["POST"])
def upload_template():
    try:
        file = request.files.get("file")
        type_notif = request.form.get("type_notification")
        intitule = request.form.get("intitule")

        if not file or not type_notif or not intitule:
            return jsonify(success=False, message="Données manquantes")

        if not file.filename.endswith(".docx"):
            return jsonify(success=False, message="Format invalide")

        folder = os.path.join("data", "templates")
        os.makedirs(folder, exist_ok=True)

        # ✅ Nom final basé sur l’intitulé
        filename = secure_filename(intitule) + ".docx"
        filepath = os.path.join(folder, filename)

        # ⚠️ Gestion doublon (important)
        if os.path.exists(filepath):
            return jsonify(success=False, message="Ce modèle existe déjà")

        file.save(filepath)

        # 📊 Mise à jour Excel
        excel_path = os.path.join(folder, "NOTIFS.xlsx")
        df = pd.read_excel(excel_path, sheet_name="NOTIFS")

        new_row = {
            "Type_notification": type_notif,
            "Modele_notification": intitule
        }

        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        df.to_excel(excel_path, sheet_name="NOTIFS", index=False)

        return jsonify(success=True, message="Modèle enregistré avec succès")

    except Exception as e:
        return jsonify(success=False, message=str(e))
        

# \modules\notification\templates\pdf.html
#-------------------------------------------------------

<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<title>Visualisation PDF</title>

<style>
/* ============================
   Reset et style global
============================ */
html, body {
    margin: 0;
    padding: 0;
    height: 100%;
    font-family: Arial, sans-serif;
    background: #f3f4f6;
}

/* ============================
   Bannière fixée en haut
============================ */
.banniere-container {
    width: 100%;
    position: fixed;
    top: 0;
    left: 0;
    z-index: 1000;
}

.banniere-container img {
    width: 100%;
    height: auto;
    display: block;
}

/* ============================
   Sidebar fixe / Menu
============================ */
.sidebar {
    position: fixed;
    top: 0; /* sera ajusté par JS */
    left: 0;
    bottom: 0;
    width: 130px;
    background: #1e293b;
    overflow-y: auto;
    padding-top: 10px;
    z-index: 900;
    transition: transform 0.3s ease;
}

.sidebar a {
    display: block;
    margin: 10px;
    padding: 10px;
    text-align: center;
    color: white;
    text-decoration: none;
    font-weight: bold;
    background-color: #127291;
    border-radius: 4px;
}

.sidebar a:hover {
    background-color: #0d5a72;
}

/* ============================
   Hamburger menu (mobile)
============================ */
.hamburger {
    display: none;
    position: fixed;
    top: 10px;
    left: 10px;
    z-index: 1100;
    cursor: pointer;
    width: 30px;
    height: 25px;
    flex-direction: column;
    justify-content: space-between;
}

.hamburger div {
    width: 100%;
    height: 4px;
    background: #127291;
    border-radius: 2px;
}

/* ============================
   Contenu principal (PDF)
============================ */
.main-content {
    margin-left: 130px; /* espace pour sidebar */
    padding: 0;
    padding-top: 0; /* ajusté par JS selon bannière */
    height: 100vh;
    box-sizing: border-box;
}

.main-content iframe {
    width: 100%;
    border: none;
    display: block;
}

/* ============================
   Responsive mobile
============================ */
@media (max-width: 768px) {
    .sidebar {
        width: 200px;
        transform: translateX(-220px);
    }
    .sidebar.active {
        transform: translateX(0);
    }
    .main-content {
        margin-left: 0;
    }
    .hamburger {
        display: flex;
    }
}
</style>
</head>
<body>

<!-- Bannière -->
<div class="banniere-container">
    <img id="banniere-img" src="{{ url_for('static', filename='images/notification.jpg') }}" alt="Bannière">
</div>

<!-- Hamburger menu -->
<div class="hamburger" id="hamburger">
    <div></div>
    <div></div>
    <div></div>
</div>

<!-- Sidebar -->
<aside class="sidebar" id="sidebar">
    <a href="/">Accueil</a>
    <a href="/notification">Restitution</a>
</aside>

<!-- Contenu principal -->
<div class="main-content" id="main-content">
    <iframe id="pdf-frame" src="/notification/pdf_file/{{ fichier }}"></iframe>
</div>

<script>
// ============================
// Ajustement dynamique layout
// ============================
function ajusterLayout() {
    const banniere = document.getElementById('banniere-img');
    const sidebar = document.getElementById('sidebar');
    const mainContent = document.getElementById('main-content');
    const iframe = document.getElementById('pdf-frame');

    const banniereHeight = banniere.clientHeight;

    sidebar.style.top = banniereHeight + 'px';
    mainContent.style.paddingTop = banniereHeight + 'px';
    iframe.style.height = `calc(100vh - ${banniereHeight}px)`;
}

window.addEventListener('load', ajusterLayout);
window.addEventListener('resize', ajusterLayout);

// ============================
// Hamburger toggle
// ============================
const hamburger = document.getElementById('hamburger');
const sidebar = document.getElementById('sidebar');

hamburger.addEventListener('click', () => {
    sidebar.classList.toggle('active');
});
</script>

</body>
</html>
       

# \modules\notification\scripts\script_notification.py
#-------------------------------------------------------
# ce script s'execute par un bouton de commande sur la page notification.html

import os
import io
import pandas as pd
from datetime import datetime
import sqlite3
from docx import Document
from docx.enum.text import WD_BREAK
from copy import deepcopy
from docxcompose.composer import Composer
import subprocess
import platform
import time


# ===============================
# Fonctions utilitaires
# ===============================
def format_nombre(valeur):
    try:
        if pd.isna(valeur):
            return ""
        return "{:,.0f}".format(float(valeur)).replace(",", " ")
    except Exception:
        return str(valeur)

def format_taux(valeur):
    try:
        if pd.isna(valeur) or valeur == "":
            return ""
        return "{:,.2f}".format(float(valeur)).replace(",", " ").replace(".", ",")
    except Exception:
        return str(valeur)

# ===============================
# Script principal
# ===============================
def run_notif(
        fichier,
        niu,
        centre,
        cri,
        type_personne,
        type_notification,
        modele,
        sortie_dir,
        db_path,
        utilisateur
):
    start_time = time.time()
    print("[INFO] Démarrage du script de notification.")
    print("UTILISATEUR RECU :", utilisateur)

    BASE_DIR = os.getcwd()  # pour construire les chemins absolus

    # 1️⃣ Chargement fichier Excel
    input_path = os.path.join(sortie_dir, fichier)
    if not os.path.exists(input_path):
        return {"message": "Fichier introuvable."}

    df = pd.read_excel(input_path, engine="openpyxl")
    df.columns = df.columns.str.strip()
    if df.empty:
        return {"message": "Fichier Excel vide."}
    print(f"[INFO] Fichier Excel chargé avec {len(df)} lignes.")

    # 2️⃣ Filtrage
    if niu:
        df = df[df["NIU"] == niu]
    else:
        if type_personne:
            df = df[df["NIU"].astype(str).str[0] == type_personne]
        if cri:
            df = df[df["ACRI"] == cri]
        if centre:
            df = df[df["CENTRE"] == centre]
    if df.empty:
        return {"message": "Aucune donnée trouvée après filtrage."}

    # 3️⃣ Construction nom Excel filtré
    nouveau_nom = fichier.replace("Output_", "_")
    nom_excel_filtre = os.path.splitext(nouveau_nom)[0] + ".xlsx"

    # 4️⃣ Filtre dynamique
    elements_filtre = [str(x) for x in [niu, type_personne, cri, centre] if x]
    filtre = "_".join(elements_filtre) if elements_filtre else "GLOBAL"
    for c in ['<', '>', ':', '"', '/', '\\', '|', '?']:
        filtre = filtre.replace(c, "-")
    filtre = filtre.replace(" ", "_")
    print(f"[INFO] Filtre dynamique généré : {filtre}")

    # 5️⃣ Chargement modèle Word
    modele = modele.strip()
    if not modele.lower().endswith(".docx"):
        modele += ".docx"
    template_path = os.path.join(BASE_DIR, "data", "Templates", modele)
    if not os.path.exists(template_path):
        return {"message": f"❌ Modèle introuvable : {template_path}"}
    print("[INFO] Modèle Word chargé :", template_path)
    base_doc = Document(template_path)

    # 6️⃣ Fonctions de remplacement
    def remplacer_bloc(paragraphes, mapping):
        for p in paragraphes:
            full_text = "".join(run.text for run in p.runs)
            if not full_text: continue
            for key, value in mapping.items():
                if key in full_text: full_text = full_text.replace(key, str(value))
            if p.runs:
                p.runs[0].text = full_text
                for run in p.runs[1:]: run.text = ""       

    def remplacer_doc(doc, mapping):
        remplacer_bloc(doc.paragraphs, mapping)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    remplacer_bloc(cell.paragraphs, mapping)
        for section in doc.sections:
            remplacer_bloc(section.header.paragraphs, mapping)
            remplacer_bloc(section.footer.paragraphs, mapping)
            if section.first_page_header:
                remplacer_bloc(section.first_page_header.paragraphs, mapping)
            if section.first_page_footer:
                remplacer_bloc(section.first_page_footer.paragraphs, mapping)
            if section.even_page_header:
                remplacer_bloc(section.even_page_header.paragraphs, mapping)
            if section.even_page_footer:
                remplacer_bloc(section.even_page_footer.paragraphs, mapping)

    # 7️⃣ Génération bulletins individuels en mémoire
    date_ref = datetime.now().strftime("%Y%m%d")
    bulletins_mem = []

    for index, row in df.iterrows():
        doc = deepcopy(base_doc)
        mapping = {
            "{{NOTIF}}": type_notification,
            "{{REF_NOTIF}}": f"UTTAD-{index+1}-{date_ref}",
            "{{DIRECTEUR}}": row.get("DIRECTEUR", ""),
            "{{CRIF}}": row.get("CRIF", ""),
            "{{CDIF}}": row.get("CDIF", ""),
            "{{CRIA}}": row.get("CRIA", ""),
            "{{CDIA}}": row.get("CDIA", ""),
            "{{ACRI}}": row.get("ACRI", ""),
            "{{ACDI}}": row.get("ACDI", ""),
            "{{NIU}}": row.get("NIU", ""),
            "{{NOM_RAISON}}": row.get("NOM_RAISON", ""),
            "{{SIGLE}}": row.get("SIGLE", ""),
            "{{TELEPHONE}}": format_nombre(row.get("TELEPHONE")),
            "{{ACTIVITE}}": row.get("ACTIVITE", ""),
            "{{TITRE}}": row.get("TITRE", ""),
            "{{APPEL}}": row.get("APPEL", ""),
            "{{VILLE}}": row.get("VILLE", ""),
            "{{AXE}}": row.get("AXE", ""),
            "{{MATIERE}}": row.get("MATIERE", ""),
            "{{VAL_DECLA}}": format_nombre(row.get("VAL_DECLA")),
            "{{VAL_RECOUP}}": format_nombre(row.get("VAL_RECOUP")),
            "{{ECART_NOTIF}}": format_nombre(row.get("ECART_NOTIF")),
            "{{MARGE_TAUX}}": format_nombre(row.get("MARGE_TAUX")),
            "{{MARGE}}": format_nombre(row.get("MARGE")),
            "{{BASE}}": format_nombre(row.get("BASE")),
            "{{PERIODE}}": row.get("PERIODE", ""),
            "{{LISTE_IMPOTS}}": row.get("LISTE_IMPOTS", ""),
            "{{AUTAX}}": row.get("AUTRE_TAXE", ""),
            "{{PROVENANCE_A}}": row.get("PROVENANCE_A", ""),
            "{{PROVENANCE_B}}": row.get("PROVENANCE_B", ""),
            "{{SOURCE_A}}": row.get("SOURCE_A", ""),
            "{{SOURCE_B}}": row.get("SOURCE_B", ""),
            "{{PERIODE_A}}": row.get("PERIODE_A", ""),
            "{{PERIODE_B}}": row.get("PERIODE_B", ""),
            "{{IS_BASE}}": format_nombre(row.get("IS_BASE")),
            "{{IS_TAUX}}": format_taux(row.get("IS_TAUX")),
            "{{IS_PRINCI}}": format_nombre(row.get("IS_PRINCI")),
            "{{IS_PENAL}}": format_nombre(row.get("IS_PENAL")),
            "{{IS_TOTAL}}": format_nombre(row.get("IS_TOTAL")),
            "{{IRPP_BASE}}": format_nombre(row.get("IRPP_BASE")),
            "{{IRPP_TAUX}}": format_taux(row.get("IRPP_TAUX")),
            "{{IRPP_PRINCI}}": format_nombre(row.get("IRPP_PRINCI")),
            "{{IRPP_PENAL}}": format_nombre(row.get("IRPP_PENAL")),
            "{{IRPP_TOTAL}}": format_nombre(row.get("IRPP_TOTAL")), 
            "{{IRCM_BASE}}": format_nombre(row.get("IRCM_BASE")),
            "{{IRCM_TAUX}}": format_taux(row.get("IRCM_TAUX")),
            "{{IRCM_PRINCI}}": format_nombre(row.get("IRCM_PRINCI")),
            "{{IRCM_PENAL}}": format_nombre(row.get("IRCM_PENAL")),
            "{{IRCM_TOTAL}}": format_nombre(row.get("IRCM_TOTAL")),
            "{{TVA_BASE}}": format_nombre(row.get("TVA_BASE")),
            "{{TVA_TAUX}}": format_taux(row.get("TVA_TAUX")),
            "{{TVA_PRINCI}}": format_nombre(row.get("TVA_PRINCI")),
            "{{TVA_PENAL}}": format_nombre(row.get("TVA_PENAL")),
            "{{TVA_TOTAL}}": format_nombre(row.get("TVA_TOTAL")),
            "{{TAX_PRINCI}}": format_nombre(row.get("TAX_PRINCI")),
            "{{TAX_PENAL}}": format_nombre(row.get("TAX_PENAL")),
            "{{TAX_TOTAL}}": format_nombre(row.get("TAX_TOTAL")),
            "{{AUT_BASE}}": format_nombre(row.get(f"{row.get("AUTRE_TAXE", "")}_BASE")),
            "{{AUT_TAUX}}": format_taux(row.get(f"{row.get("AUTRE_TAXE", "")}_TAUX")),
            "{{AUT_PRINCI}}": format_nombre(row.get(f"{row.get("AUTRE_TAXE", "")}_PRINCI")),
            "{{AUT_PENAL}}": format_nombre(row.get(f"{row.get("AUTRE_TAXE", "")}_PENAL")),
            "{{AUT_TOTAL}}": format_nombre(row.get(f"{row.get("AUTRE_TAXE", "")}_TOTAL"))
        }
        remplacer_doc(doc, mapping)

        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        bulletins_mem.append(mem_file)

    print(f"[INFO] {len(bulletins_mem)} bulletins générés en mémoire.")

    # 8️⃣ Fusion Word finale
    nom_doc = f"BI{os.path.splitext(nouveau_nom)[0]}_{filtre}.docx"
    nom_docx_final = os.path.join(sortie_dir, nom_doc)

    if bulletins_mem:
        master_doc = Document(bulletins_mem[0])
        composer = Composer(master_doc)

        for mem_file in bulletins_mem[1:]:
            # 🔹 Ajouter un saut de page
            master_doc.add_page_break()

            # 🔹 Ajouter le document suivant
            composer.append(Document(mem_file))

        composer.save(nom_docx_final)
        print(f"[INFO] Fusion Word finale sauvegardée : {nom_docx_final}")

    # 9️⃣ Conversion PDF
    nom_pdf = nom_doc.replace(".docx", ".pdf")
    chemin_pdf = os.path.join(sortie_dir, nom_pdf)
    try:
        system_os = platform.system()
        libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe" if system_os == "Windows" else "soffice"
        subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', sortie_dir, nom_docx_final],
                       capture_output=True, text=True)
        if not os.path.exists(chemin_pdf):
            return {"message": "La conversion PDF a échoué."}
        print(f"[INFO] Conversion PDF réussie : {chemin_pdf}")
    except Exception as e:
        return {"message": f"Exception lors de la conversion PDF : {e}"}

    # 1️⃣1️⃣ Enregistrement base SQLite
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO notification (
            notification,
            nom_fichier,
            type_notification,
            script_utilise,
            utilisateur,
            date_creation
        )
        VALUES (?, ?, ?, ?, ?, ?)
    """, (
        f"Notification {nom_excel_filtre}",
        nom_pdf,
        type_notification,
        type_notification,
        utilisateur,
        datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    ))
    conn.commit()
    conn.close()

    end_time = time.time()
    duree = end_time - start_time
    print(f"[INFO] Script terminé avec succès en {duree:.2f} secondes.")

    return {"message": f"{len(df)} bulletins générés avec succès en {duree:.2f} secondes.", "pdf_final": nom_pdf}