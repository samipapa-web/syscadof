# \modules\notification\scripts\script_notification.py
#-------------------------------------------------------

import os
import time
import pandas as pd
from datetime import datetime
from docx import Document
from copy import deepcopy
from docxcompose.composer import Composer
from flask import current_app
from werkzeug.utils import secure_filename
from psycopg2.extras import RealDictCursor

# =========================================================
# HISTORIQUE
# =========================================================
def save_historique(conn, notification, nom_fichier, type_notification, utilisateur):
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute("""
            INSERT INTO notification (
                notification,
                nom_fichier,
                type_notification,
                utilisateur,
                date_creation
            ) VALUES (%s, %s, %s, %s, NOW())
        """, (notification, nom_fichier, type_notification, utilisateur))


# =========================================================
# FORMATAGE
# =========================================================
def format_nombre(valeur):
    try:
        if pd.isna(valeur) or valeur == "":
            return ""
        return "{:,.0f}".format(float(valeur)).replace(",", " ")
    except:
        return str(valeur)


def format_taux(valeur):
    try:
        if pd.isna(valeur) or valeur == "":
            return ""
        return "{:,.2f}".format(float(valeur)).replace(",", " ").replace(".", ",")
    except:
        return str(valeur)


# =========================================================
# RUN NOTIFICATION (OPTIMISÉ)
# =========================================================
def run_notif(
    fichier,
    niu,
    centre,
    cri,
    type_personne,
    type_notification,
    modele,
    sortie_dir,
    conn,
    utilisateur
):

    start_time = time.time()

    fichier = secure_filename(fichier)

    input_path = os.path.join(sortie_dir, fichier)
    if not os.path.exists(input_path):
        return {"message": "Fichier introuvable"}

    # =====================================================
    # 1. CHARGEMENT DATA
    # =====================================================
    df = pd.read_excel(input_path, engine="openpyxl", dtype=str)
    df.columns = df.columns.str.strip()

    if df.empty:
        return {"message": "Fichier vide"}

    # =====================================================
    # 2. FILTRAGE
    # =====================================================
    if niu:
        df = df[df["NIU"] == str(niu)]
    else:
        if type_personne:
            df = df[df["NIU"].astype(str).str[0] == type_personne]
        if cri:
            df = df[df["ACRI"] == cri]
        if centre:
            df = df[df["CENTRE"] == centre]

    if df.empty:
        return {"message": "Aucune donnée après filtrage"}

    # =====================================================
    # 3. MODELE
    # =====================================================
    modele = secure_filename(modele.strip())
    if not modele.endswith(".docx"):
        modele += ".docx"

    TEMPLATES_DIR = current_app.config.get("TEMPLATES_DIR")
    template_path = os.path.join(TEMPLATES_DIR, modele)

    if not os.path.exists(template_path):
        return {"message": "Modèle introuvable"}

    # =====================================================
    # 4. FONCTIONS DE PUBLIPOSTAGE
    # =====================================================

    def remplacer_runs(paragraph, mapping):
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)

        for key, value in mapping.items():
            if key in full_text:

                # 1. tentative remplacement simple (run par run)
                replaced = False
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))
                        replaced = True

                # 2. fallback si variable éclatée sur plusieurs runs
                if not replaced:
                    new_text = full_text.replace(key, str(value))

                    # ⚠️ on ne recrée PAS brutalement tous les runs
                    # on conserve le premier run (style principal)
                    paragraph.runs[0].text = new_text

                    for run in paragraph.runs[1:]:
                        run.text = ""
            
    def remplacer_doc_complet(doc, mapping):

        def traiter_paragraphes(paragraphs):
            for p in paragraphs:
                remplacer_runs(p, mapping)

        # Paragraphes principaux
        traiter_paragraphes(doc.paragraphs)

        # Tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    traiter_paragraphes(cell.paragraphs)

        # Sections (headers/footers)
        for section in doc.sections:

            traiter_paragraphes(section.header.paragraphs)
            traiter_paragraphes(section.footer.paragraphs)

            if section.first_page_header:
                traiter_paragraphes(section.first_page_header.paragraphs)

            if section.first_page_footer:
                traiter_paragraphes(section.first_page_footer.paragraphs)

            if section.even_page_header:
                traiter_paragraphes(section.even_page_header.paragraphs)

            if section.even_page_footer:
                traiter_paragraphes(section.even_page_footer.paragraphs)

    # =====================================================
    # 5. PREPARATION FUSION (DISQUE AU LIEU RAM)
    # =====================================================
    
    date_ref = datetime.now().strftime("%Y%m%d")
        
    temp_files = []

    base_doc = Document(template_path)

    for index, row in df.iterrows():

        doc = deepcopy(base_doc)

        autre_taxe = str(row.get("AUTRE_TAXE") or "")

        mapping = {
            "{{NOTIF}}": type_notification,
            "{{REF_NOTIF}}": f"UTTAD-{index+1}-{date_ref}",
            "{{DIRECTEUR}}": row.get("DIRECTEUR", ""),
            "{{CRIF}}": row.get("CRIF", ""),
            "{{CDIF}}": row.get("CDIF", ""),
            "{{CRIA}}": row.get("CRIA", ""),
            "{{CDIA}}": row.get("CDIA", ""),
            "{{ACRI}}": row.get("ACRI", ""),
            "{{ACDI}}": row.get("ACDI", ""),
            "{{NIU}}": row.get("NIU", ""),
            "{{NOM_RAISON}}": row.get("NOM_RAISON", ""),
            "{{SIGLE}}": row.get("SIGLE", ""),
            "{{TELEPHONE}}": format_nombre(row.get("TELEPHONE")),
            "{{ACTIVITE}}": row.get("ACTIVITE", ""),
            "{{TITRE}}": row.get("TITRE", ""),
            "{{APPEL}}": row.get("APPEL", ""),
            "{{VILLE}}": row.get("VILLE", ""),
            "{{AXE}}": row.get("AXE", ""),
            "{{MATIERE}}": row.get("MATIERE", ""),
            "{{VAL_DECLA}}": format_nombre(row.get("VAL_DECLA")),
            "{{VAL_RECOUP}}": format_nombre(row.get("VAL_RECOUP")),
            "{{ECART_NOTIF}}": format_nombre(row.get("ECART_NOTIF")),
            "{{MARGE_TAUX}}": format_nombre(row.get("MARGE_TAUX")),
            "{{MARGE}}": format_nombre(row.get("MARGE")),
            "{{BASE}}": format_nombre(row.get("BASE")),
            "{{PERIODE}}": row.get("PERIODE", ""),
            "{{LISTE_IMPOTS}}": row.get("LISTE_IMPOTS", ""),
            "{{AUTAX}}": row.get("AUTRE_TAXE", ""),
            "{{PROVENANCE_A}}": row.get("PROVENANCE_A", ""),
            "{{PROVENANCE_B}}": row.get("PROVENANCE_B", ""),
            "{{SOURCE_A}}": row.get("SOURCE_A", ""),
            "{{SOURCE_B}}": row.get("SOURCE_B", ""),
            "{{PERIODE_A}}": row.get("PERIODE_A", ""),
            "{{PERIODE_B}}": row.get("PERIODE_B", ""),
            "{{IS_BASE}}": format_nombre(row.get("IS_BASE")),
            "{{IS_TAUX}}": format_taux(row.get("IS_TAUX")),
            "{{IS_PRINCI}}": format_nombre(row.get("IS_PRINCI")),
            "{{IS_PENAL}}": format_nombre(row.get("IS_PENAL")),
            "{{IS_TOTAL}}": format_nombre(row.get("IS_TOTAL")),
            "{{IRPP_BASE}}": format_nombre(row.get("IRPP_BASE")),
            "{{IRPP_TAUX}}": format_taux(row.get("IRPP_TAUX")),
            "{{IRPP_PRINCI}}": format_nombre(row.get("IRPP_PRINCI")),
            "{{IRPP_PENAL}}": format_nombre(row.get("IRPP_PENAL")),
            "{{IRPP_TOTAL}}": format_nombre(row.get("IRPP_TOTAL")), 
            "{{IRCM_BASE}}": format_nombre(row.get("IRCM_BASE")),
            "{{IRCM_TAUX}}": format_taux(row.get("IRCM_TAUX")),
            "{{IRCM_PRINCI}}": format_nombre(row.get("IRCM_PRINCI")),
            "{{IRCM_PENAL}}": format_nombre(row.get("IRCM_PENAL")),
            "{{IRCM_TOTAL}}": format_nombre(row.get("IRCM_TOTAL")),
            "{{TVA_BASE}}": format_nombre(row.get("TVA_BASE")),
            "{{TVA_TAUX}}": format_taux(row.get("TVA_TAUX")),
            "{{TVA_PRINCI}}": format_nombre(row.get("TVA_PRINCI")),
            "{{TVA_PENAL}}": format_nombre(row.get("TVA_PENAL")),
            "{{TVA_TOTAL}}": format_nombre(row.get("TVA_TOTAL")),
            "{{TAX_PRINCI}}": format_nombre(row.get("TAX_PRINCI")),
            "{{TAX_PENAL}}": format_nombre(row.get("TAX_PENAL")),
            "{{TAX_TOTAL}}": format_nombre(row.get("TAX_TOTAL")),
            "{{AUT_BASE}}": format_nombre(row.get(f"{autre_taxe}_BASE")),
            "{{AUT_TAUX}}": format_taux(row.get(f"{autre_taxe}_TAUX")),
            "{{AUT_PRINCI}}": format_nombre(row.get(f"{autre_taxe}_PRINCI")),
            "{{AUT_PENAL}}": format_nombre(row.get(f"{autre_taxe}_PENAL")),
            "{{AUT_TOTAL}}": format_nombre(row.get(f"{autre_taxe}_TOTAL")),
        }

        remplacer_doc_complet(doc, mapping)

        temp_path = os.path.join(
            sortie_dir,
            f"tmp_{index}_{date_ref}.docx"
        )

        doc.save(temp_path)
        temp_files.append(temp_path)

    # =====================================================
    # 6. FUSION DOCX (OPTIMISEE)
    # =====================================================
    nom_doc = f"NOTIF_{os.path.splitext(fichier)[0]}_{date_ref}.docx"
    nom_docx_final = os.path.join(sortie_dir, nom_doc)

    master_doc = Document(temp_files[0])
    composer = Composer(master_doc)

    for f in temp_files[1:]:
        composer.append(Document(f))

    composer.save(nom_docx_final)

    # =====================================================
    # 7. NETTOYAGE TEMPORAIRE
    # =====================================================
    for f in temp_files:
        if os.path.exists(f):
            os.remove(f)

    # =====================================================
    # 8. HISTORIQUE
    # =====================================================
    save_historique(
        conn=conn,
        notification=f"Notification {nom_doc}",
        nom_fichier=nom_doc,
        type_notification=type_notification,
        utilisateur=utilisateur
    )

    duree = time.time() - start_time

    return {
        "message": f"{len(df)} bulletins générés en {duree:.2f}s",
        "fichier_final": nom_doc
    }
