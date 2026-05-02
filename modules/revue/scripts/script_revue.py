# \modules\revue\scripts\script_revue.py
# -------------------------------------------------------

# \modules\revue\scripts\script_revue.py
# -------------------------------------------------------

import os
import pandas as pd
from datetime import datetime
from flask import current_app
from docx import Document
from copy import deepcopy
from docxcompose.composer import Composer


# ===============================
# FORMATAGE
# ===============================
def format_nombre(val):
    try:
        if pd.isna(val):
            return ""
        return "{:,.0f}".format(float(val)).replace(",", " ")
    except:
        return str(val)


def format_taux(val):
    try:
        if pd.isna(val):
            return ""
        return "{:,.2f}".format(float(val)).replace(",", " ").replace(".", ",")
    except:
        return str(val)


# ===============================
# SCRIPT PRINCIPAL
# ===============================
def run_revue(niu, cri, centre, sortie_dir, template_dir, cle=None):

    # =========================
    # 1. LECTURE FICHIERS
    # =========================
    fichiers = [
        f for f in os.listdir(sortie_dir)
        if f.startswith("Output_") and f.endswith(".xlsx")
    ]

    # 🔥 FILTRE CLE (NOUVEAU)
    if cle:
        cle = cle.lower()
        fichiers = [f for f in fichiers if cle in f.lower()]

    if not fichiers:
        return {"message": "Aucun fichier Output trouvé avec ce filtre"}

    dfs = []

    colonnes = [
        "NIU","VAL_RECOUP","VAL_DECLA","ECART_NOTIF","RAISON_SOCIALE","SIGLE",
        "ACTIVITE","REGIME","CENTRE","TELEPHONE","CDIF","CRIF","CDIA","CRIA",
        "ACDI","ACRI","VILLE","DIRECTEUR","REGIONAL","SOURCE_A","SOURCE_B",
        "PERIODE_A","PERIODE_B","PROVENANCE_A","PROVENANCE_B","PERIODE",
        "REF_NOTIF","RISQUE","AXE","MATIERE","LISTE_IMPOTS","AUTRE_TAXE",
        "MARGE_TAUX","MARGE","BASE","TVA_BASE","TVA_TAUX","TVA_PRINCI",
        "TVA_PENAL","TVA_TOTAL","IS_BASE","IS_TAUX","IS_PRINCI","IS_PENAL",
        "IS_TOTAL","IRPP_BASE","IRPP_TAUX","IRPP_PRINCI","IRPP_PENAL",
        "IRPP_TOTAL","IRCM_BASE","IRCM_TAUX","IRCM_PRINCI","IRCM_PENAL",
        "IRCM_TOTAL","BB_BASE","BB_TAUX","BB_PRINCI","BB_PENAL","BB_TOTAL",
        "TAX_PRINCI","TAX_PENAL","TAX_TOTAL","TITRE","APPEL"
    ]

    # =========================
    # 2. FILTRAGE + CONSOLIDATION
    # =========================
    for f in fichiers:
        path = os.path.join(sortie_dir, f)

        try:
            df = pd.read_excel(path, engine="openpyxl")
        except Exception:
            continue

        df.columns = df.columns.str.strip()

        if niu:
            df = df[df["NIU"] == niu]
        if cri:
            df = df[df["ACRI"] == cri]
        if centre:
            df = df[df["CENTRE"] == centre]

        if not df.empty:
            df = df.reindex(columns=colonnes)
            df.insert(0, "FICHIER", f)
            dfs.append(df)

    if not dfs:
        return {"message": "Aucune donnée trouvée après filtrage"}

    df_final = pd.concat(dfs, ignore_index=True)

    # =========================
    # 🔥 2B. SYNTHESE DETAILLEE
    # =========================
    df_final["ECART_NOTIF"] = pd.to_numeric(df_final["ECART_NOTIF"], errors="coerce").fillna(0)
    df_final["TAX_PRINCI"] = pd.to_numeric(df_final["TAX_PRINCI"], errors="coerce").fillna(0)
    df_final["TAX_PENAL"] = pd.to_numeric(df_final["TAX_PENAL"], errors="coerce").fillna(0)
    df_final["TAX_TOTAL"] = pd.to_numeric(df_final["TAX_TOTAL"], errors="coerce").fillna(0)

    synthese_par_fichier = (
        df_final
        .groupby("FICHIER")
        .agg(
            lignes=("FICHIER", "count"),
            ecart=("ECART_NOTIF", "sum"),
            tax_princi=("TAX_PRINCI", "sum"),
            tax_penal=("TAX_PENAL", "sum"),
            tax_total=("TAX_TOTAL", "sum")
        )
        .reset_index()
        .to_dict(orient="records")
    )

    synthese = {
        "total_lignes": len(df_final),
        "total_fichiers": df_final["FICHIER"].nunique(),
        "total_ecart": float(df_final["ECART_NOTIF"].sum()),
        "total_tax_princi": float(df_final["TAX_PRINCI"].sum()),
        "total_tax_penal": float(df_final["TAX_PENAL"].sum()),
        "total_tax_total": float(df_final["TAX_TOTAL"].sum()),
    }

    # =========================
    # 3. EXPORT EXCEL (INCHANGE)
    # =========================
    date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
    excel_name = f"REVUE_{date_str}.xlsx"
    excel_path = os.path.join(sortie_dir, excel_name)

    df_final.to_excel(excel_path, index=False)

    # =========================
    # 4 → 8 STRICTEMENT INCHANGEES
    # =========================
    template_path = os.path.join(template_dir, "Rev_Sans_annexe.docx")

    if not os.path.exists(template_path):
        return {"message": "❌ Modèle revue.docx introuvable"}

    base_doc = Document(template_path)

    def remplacer_runs(paragraph, mapping):
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)

        for key, value in mapping.items():
            if key in full_text:

                replaced = False
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))
                        replaced = True

                if not replaced:
                    new_text = full_text.replace(key, str(value))
                    paragraph.runs[0].text = new_text
                    for run in paragraph.runs[1:]:
                        run.text = ""

    def remplacer_doc(doc, mapping):
        def traiter(paragraphs):
            for p in paragraphs:
                remplacer_runs(p, mapping)

        traiter(doc.paragraphs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    traiter(cell.paragraphs)

        for section in doc.sections:
            traiter(section.header.paragraphs)
            traiter(section.footer.paragraphs)

    date_ref = datetime.now().strftime("%Y%m%d")
    temp_files = []

    for index, row in df_final.iterrows():

        doc = deepcopy(base_doc)

        autre_taxe = str(row.get("AUTRE_TAXE") or "").strip()

        mapping = {
            "{{NOTIF}}": "REVUE",
            "{{REF_NOTIF}}": f"REV-{index+1}-{date_ref}",
            "{{DIRECTEUR}}": row.get("DIRECTEUR", ""),
            "{{CRIF}}": row.get("CRIF", ""),
            "{{CDIF}}": row.get("CDIF", ""),
            "{{CRIA}}": row.get("CRIA", ""),
            "{{CDIA}}": row.get("CDIA", ""),
            "{{ACRI}}": row.get("ACRI", ""),
            "{{ACDI}}": row.get("ACDI", ""),
            "{{NIU}}": row.get("NIU", ""),
            "{{NOM_RAISON}}": row.get("RAISON_SOCIALE", ""),
            "{{SIGLE}}": row.get("SIGLE", ""),
            "{{TELEPHONE}}": format_nombre(row.get("TELEPHONE")),
            "{{ACTIVITE}}": row.get("ACTIVITE", ""),
            "{{TITRE}}": row.get("TITRE", ""),
            "{{APPEL}}": row.get("APPEL", ""),
            "{{VILLE}}": row.get("VILLE", ""),
            "{{AXE}}": row.get("AXE", ""),
            "{{MATIERE}}": row.get("MATIERE", ""),
            "{{VAL_DECLA}}": format_nombre(row.get("VAL_DECLA")),
            "{{VAL_RECOUP}}": format_nombre(row.get("VAL_RECOUP")),
            "{{ECART_NOTIF}}": format_nombre(row.get("ECART_NOTIF")),
            "{{MARGE_TAUX}}": format_nombre(row.get("MARGE_TAUX")),
            "{{MARGE}}": format_nombre(row.get("MARGE")),
            "{{BASE}}": format_nombre(row.get("BASE")),
            "{{PERIODE}}": row.get("PERIODE", ""),
            "{{LISTE_IMPOTS}}": row.get("LISTE_IMPOTS", ""),
            "{{AUTAX}}": autre_taxe,
            "{{PROVENANCE_A}}": row.get("PROVENANCE_A", ""),
            "{{PROVENANCE_B}}": row.get("PROVENANCE_B", ""),
            "{{SOURCE_A}}": row.get("SOURCE_A", ""),
            "{{SOURCE_B}}": row.get("SOURCE_B", ""),
            "{{PERIODE_A}}": row.get("PERIODE_A", ""),
            "{{PERIODE_B}}": row.get("PERIODE_B", ""),
            "{{IS_BASE}}": format_nombre(row.get("IS_BASE")),
            "{{IS_TAUX}}": format_taux(row.get("IS_TAUX")),
            "{{IS_PRINCI}}": format_nombre(row.get("IS_PRINCI")),
            "{{IS_PENAL}}": format_nombre(row.get("IS_PENAL")),
            "{{IS_TOTAL}}": format_nombre(row.get("IS_TOTAL")),
            "{{IRPP_BASE}}": format_nombre(row.get("IRPP_BASE")),
            "{{IRPP_TAUX}}": format_taux(row.get("IRPP_TAUX")),
            "{{IRPP_PRINCI}}": format_nombre(row.get("IRPP_PRINCI")),
            "{{IRPP_PENAL}}": format_nombre(row.get("IRPP_PENAL")),
            "{{IRPP_TOTAL}}": format_nombre(row.get("IRPP_TOTAL")),
            "{{IRCM_BASE}}": format_nombre(row.get("IRCM_BASE")),
            "{{IRCM_TAUX}}": format_taux(row.get("IRCM_TAUX")),
            "{{IRCM_PRINCI}}": format_nombre(row.get("IRCM_PRINCI")),
            "{{IRCM_PENAL}}": format_nombre(row.get("IRCM_PENAL")),
            "{{IRCM_TOTAL}}": format_nombre(row.get("IRCM_TOTAL")),
            "{{TVA_BASE}}": format_nombre(row.get("TVA_BASE")),
            "{{TVA_TAUX}}": format_taux(row.get("TVA_TAUX")),
            "{{TVA_PRINCI}}": format_nombre(row.get("TVA_PRINCI")),
            "{{TVA_PENAL}}": format_nombre(row.get("TVA_PENAL")),
            "{{TVA_TOTAL}}": format_nombre(row.get("TVA_TOTAL")),
            "{{TAX_PRINCI}}": format_nombre(row.get("TAX_PRINCI")),
            "{{TAX_PENAL}}": format_nombre(row.get("TAX_PENAL")),
            "{{TAX_TOTAL}}": format_nombre(row.get("TAX_TOTAL")),
        }

        if autre_taxe:
            mapping.update({
                "{{AUT_BASE}}": format_nombre(row.get(f"{autre_taxe}_BASE")),
                "{{AUT_TAUX}}": format_taux(row.get(f"{autre_taxe}_TAUX")),
                "{{AUT_PRINCI}}": format_nombre(row.get(f"{autre_taxe}_PRINCI")),
                "{{AUT_PENAL}}": format_nombre(row.get(f"{autre_taxe}_PENAL")),
                "{{AUT_TOTAL}}": format_nombre(row.get(f"{autre_taxe}_TOTAL")),
            })

        remplacer_doc(doc, mapping)

        temp_path = os.path.join(sortie_dir, f"tmp_{index}_{date_ref}.docx")
        doc.save(temp_path)
        temp_files.append(temp_path)

    docx_name = f"REVUE_{date_str}.docx"
    docx_path = os.path.join(sortie_dir, docx_name)

    master_doc = Document(temp_files[0])
    composer = Composer(master_doc)

    for f in temp_files[1:]:
        master_doc.add_page_break()
        composer.append(Document(f))

    composer.save(docx_path)

    for f in temp_files:
        if os.path.exists(f):
            os.remove(f)

    # =========================
    # 9. RESULTAT FINAL ENRICHI
    # =========================
    return {
        "message": f"{len(df_final)} lignes consolidées",
        "excel": excel_name,
        "docx": docx_name,
        "synthese": synthese,
        "detail_par_fichier": synthese_par_fichier
    }
